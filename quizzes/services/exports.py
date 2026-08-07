import csv
import io
import random
import re

from django.http import HttpResponse
from docx import Document

from . import sampling
from .grading import ensure_not_expired

DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def _safe_filename(name: str) -> str:
    return re.sub(r"[^A-Za-z0-9_.-]+", "_", name)


def export_attempts_csv(quiz) -> HttpResponse:
    # Self-correct any attempts nobody has revisited past their deadline,
    # so exports are always accurate at the moment they're needed.
    for attempt in quiz.attempts.filter(submitted_at__isnull=True):
        ensure_not_expired(attempt)

    response = HttpResponse(content_type="text/csv")
    filename = _safe_filename(f"{quiz.title}_results.csv")
    response["Content-Disposition"] = f'attachment; filename="{filename}"'

    writer = csv.writer(response)
    writer.writerow(
        [
            "username",
            "first_name",
            "last_name",
            "student_email",
            "score_percent",
            "passed",
            "submitted_at",
            "auto_submitted",
        ]
    )
    attempts = (
        quiz.attempts.filter(submitted_at__isnull=False)
        .select_related("student")
        .order_by("student__last_name", "student__first_name", "submitted_at")
    )
    for attempt in attempts:
        writer.writerow(
            [
                attempt.student.username,
                attempt.student.first_name,
                attempt.student.last_name,
                attempt.student.email,
                attempt.score_percent,
                attempt.passed,
                attempt.submitted_at.isoformat(),
                attempt.auto_submitted,
            ]
        )
    return response


def _render_quiz_docx(title: str, question_rows) -> io.BytesIO:
    """question_rows: iterable of (position, question_text, options)."""
    document = Document()
    document.add_heading(title, level=1)
    for position, question_text, options in question_rows:
        document.add_paragraph(f"{position}. {question_text}")
        for opt in options:
            document.add_paragraph(f"    {opt['label']}. {opt['text']}")
        document.add_paragraph("")
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def export_attempt_docx(attempt) -> HttpResponse:
    """Export the frozen questions actually served for a specific attempt --
    what the student saw, in the order and lettering they saw it."""
    question_rows = [
        (aq.position, aq.bank_item.question_text, aq.shuffled_options)
        for aq in attempt.questions.select_related("bank_item").order_by("position")
    ]
    buffer = _render_quiz_docx(attempt.quiz.title, question_rows)
    response = HttpResponse(buffer.read(), content_type=DOCX_CONTENT_TYPE)
    filename = _safe_filename(f"{attempt.quiz.title}_attempt_{attempt.id}.docx")
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


def export_blank_sample_docx(quiz) -> HttpResponse:
    """A fresh random sample for printing a paper copy of the quiz. Not tied
    to any student/attempt -- entering a paper test-taker's answers back
    into the system afterward is explicitly out of scope for v1."""
    bank_items = list(quiz.bank_items.all())
    sampled = random.sample(bank_items, k=quiz.quiz_length or len(bank_items))

    question_rows = []
    for position, bank_item in enumerate(sampled, start=1):
        shuffled_options, _ = sampling.relettered_options(
            bank_item.options, bank_item.correct_label
        )
        question_rows.append((position, bank_item.question_text, shuffled_options))

    buffer = _render_quiz_docx(quiz.title, question_rows)
    response = HttpResponse(buffer.read(), content_type=DOCX_CONTENT_TYPE)
    filename = _safe_filename(f"{quiz.title}_sample.docx")
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response
